"""Course plan builder — deterministic, re-runnable, atomic (lumen build pattern).

Turns a Phase-0 diagnosis brief + competency tree + task bank into a
concrete 4-module course plan (concept sequences + practice tasks + a
prerequisite DAG). The build is:

* **Deterministic** — same inputs ⇒ same plan, so it is safe to re-run.
* **Side-effect free** — pure computation; nothing is written until the
  caller persists via :meth:`CoursePlanStore.save`.
* **Atomic** — the plan is written with one atomic write, so an interrupted
  build never leaves a half-finished plan on disk.

The plan mirrors the coach's 4 teaching modules and is what the personal
centre / concept map can render as a "your path" view.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
import json
from pathlib import Path
from typing import Any

from deeptutor.services.learning_records import LearningRecordStore

MODULES = (
    "标注基础",
    "进阶技能",
    "质量管控",
    "工具进阶",
)

# Deterministic module→task rules (by type + difficulty).
_TASK_RULES: tuple[tuple[str, tuple[str, ...], tuple[str, ...]], ...] = (
    ("标注基础", ("bbox",), ("easy",)),
    ("进阶技能", ("bbox",), ("medium", "hard")),
    ("质量管控", ("classification",), ("medium", "hard")),
    ("工具进阶", ("classification",), ("easy",)),
)

_TASK_MODULE = {
    "task1": "标注基础",
    "task3": "标注基础",
    "task5": "标注基础",
    "task2": "进阶技能",
    "task4": "进阶技能",
    "task6": "进阶技能",
    "task7": "进阶技能",
    "task8": "进阶技能",
    "task9": "质量管控",
}


@dataclass
class CoursePlan:
    plan_id: str
    built_at: str
    teaching_mode: str
    goal_type: str
    modules: list[dict[str, Any]]
    dag: dict[str, list[str]]
    source: str

    def to_dict(self) -> dict[str, Any]:
        return {
            "plan_id": self.plan_id,
            "built_at": self.built_at,
            "teaching_mode": self.teaching_mode,
            "goal_type": self.goal_type,
            "modules": self.modules,
            "dag": self.dag,
            "source": self.source,
        }


def _concepts_for(module: str, tree: dict) -> list[str]:
    """Deterministic concept list for a module from the competency tree."""
    group_by_id = {
        "task-1-1": "标注基础",
        "task-1-2": "标注基础",
        "task-3-1": "质量管控",
        "task-3-2": "质量管控",
        "task-3-3": "质量管控",
        "task-4-1": "工具进阶",
        "task-4-2": "工具进阶",
        "task-2-1": "进阶技能",
        "task-2-2": "进阶技能",
        "task-1-3": "进阶技能",
        "task-1-4": "进阶技能",
    }
    concepts: list[str] = []

    def walk(node: dict) -> None:
        if node.get("level") == 4:
            return
        if group_by_id.get(node.get("id")) == module and node.get("level") == 3:
            concepts.extend(
                skill.get("name", "")
                for skill in node.get("skills", [])
                if skill.get("name")
            )
        for child in node.get("children", []):
            walk(child)
        for skill in node.get("skills", []):
            walk(skill)

    walk(tree)
    return concepts


def build_course_plan(
    *,
    brief: dict[str, Any] | None = None,
    plan_id: str | None = None,
    tree: dict | None = None,
    bank: dict | None = None,
) -> CoursePlan:
    """Deterministically build a course plan. Pure — no I/O."""
    from deeptutor.tools.competency_tool import _load_competency_tree
    from deeptutor.tools.task_bank_tool import _load_bank

    tree = tree or _load_competency_tree().get("tree", {})
    bank = bank or _load_bank()
    brief = brief or {}

    teaching_mode = str(brief.get("teaching_mode") or "Standard")
    goal_type = str(brief.get("goal_type") or "interest")
    built_at = datetime.now(tz=timezone.utc).isoformat()
    plan_id = plan_id or f"plan_{built_at[:19].replace(':', '').replace('-', '')}"

    # Concept sequences (concepts appear in tree order → deterministic).
    concept_by_module: dict[str, list[str]] = {}
    for module in MODULES:
        concept_by_module[module] = _concepts_for(module, tree)

    # Tasks per module (bank insertion order → deterministic).
    task_by_module: dict[str, list[str]] = {m: [] for m in MODULES}
    for tid in sorted(bank.keys()):
        task_by_module[_TASK_MODULE.get(tid, "标注基础")].append(tid)

    # Prerequisite DAG from the tree's skill prerequisites.
    dag: dict[str, list[str]] = {}

    def walk_dag(node: dict) -> None:
        for skill in node.get("skills", []):
            prereq_ids = [p.get("id") for p in skill.get("prerequisites", [])]
            if prereq_ids:
                dag[skill.get("id", "")] = prereq_ids
        for child in node.get("children", []):
            walk_dag(child)

    walk_dag(tree)

    targets = {
        "标注基础": "F1 ≥ 0.7",
        "进阶技能": "F1 ≥ 0.85",
        "质量管控": "能自检互检",
        "工具进阶": "对标四级标准",
    }

    modules: list[dict[str, Any]] = []
    for module in MODULES:
        modules.append(
            {
                "name": module,
                "concepts": concept_by_module.get(module, []),
                "tasks": task_by_module.get(module, []),
                "target": targets.get(module, ""),
            }
        )

    return CoursePlan(
        plan_id=plan_id,
        built_at=built_at,
        teaching_mode=teaching_mode,
        goal_type=goal_type,
        modules=modules,
        dag=dag,
        source="deterministic_builder",
    )


class CoursePlanStore:
    """Persists the latest course plan atomically (re-runnable, no partials)."""

    def __init__(self) -> None:
        self._root = LearningRecordStore().file.parent
        self._file = self._root / "course_plan.json"

    @property
    def file(self) -> Path:
        return self._file

    def save(self, plan: CoursePlan) -> Path:
        self._root.mkdir(parents=True, exist_ok=True)
        from deeptutor.services.file_io import atomic_write_json

        atomic_write_json(self._file, plan.to_dict())
        return self._file

    def get(self) -> dict[str, Any] | None:
        if not self._file.exists():
            return None
        try:
            with open(self._file, encoding="utf-8") as handle:
                data = json.load(handle)
            return data if isinstance(data, dict) else None
        except (json.JSONDecodeError, OSError):
            return None

    def export_docx(self, plan: dict[str, Any] | None = None) -> dict[str, Any]:
        """Generate a readable '学习路径手册' .docx and expose it publicly.

        Returns ``{"url": ..., "path": ...}`` — a download link served under
        ``/api/outputs``. One artifact per module (concepts + tasks + target).
        """
        from deeptutor.services.path_service import get_path_service

        plan = plan or self.get() or {}
        docx = _build_docx(plan)

        try:
            task_dir = get_path_service().get_task_workspace("chat", "course_plan")
            demo_dir = task_dir / "demos"
            demo_dir.mkdir(parents=True, exist_ok=True)
            out = demo_dir / "learning_path.docx"
            docx.save(str(out))

            from deeptutor.services.sandbox.artifacts import collect_public_artifacts

            artifacts = collect_public_artifacts(str(demo_dir))
            if not artifacts:
                return {"url": None, "path": str(out)}
            return {"url": artifacts[0].url, "path": str(out)}
        except Exception:
            return {"url": None, "path": str(out)}


def rebuild(force: bool = False) -> dict[str, Any] | None:
    """Rebuild and persist the course plan from the latest brief.

    ``force=False`` reuses an existing plan (idempotent re-run); ``force=True``
    rebuilds from scratch. Returns the persisted plan dict.
    """
    store = CoursePlanStore()
    existing = store.get()
    if existing and not force:
        return existing
    brief = LearningRecordStore().get_brief()
    plan = build_course_plan(brief=brief)
    store.save(plan)
    return plan.to_dict()


def _build_docx(plan: dict[str, Any]):
    """Build a readable 学习路径手册 .docx from a course-plan dict."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading("AI 数据标注工程师 — 学习路径手册", level=0)
    doc.add_paragraph(
        f"教学模式: {plan.get('teaching_mode', 'Standard')}  |  "
        f"学习目标: {plan.get('goal_type', 'interest')}"
    )
    doc.add_paragraph(f"计划 ID: {plan.get('plan_id', '')}  |  生成: {plan.get('built_at', '')[:16]}")
    doc.add_paragraph("")

    for module in plan.get("modules", []):
        doc.add_heading(module.get("name", ""), level=1)
        doc.add_paragraph(f"目标: {module.get('target', '')}")
        concepts = module.get("concepts") or []
        tasks = module.get("tasks") or []
        if concepts:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("概念序列: ").bold = True
            p.add_run(" → ".join(concepts))
        if tasks:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("练习任务: ").bold = True
            p.add_run(", ".join(tasks))
        doc.add_paragraph("")

    dag = plan.get("dag") or {}
    if dag:
        doc.add_heading("技能前置关系", level=1)
        for k, v in list(dag.items())[:15]:
            doc.add_paragraph(f"{k} ← {'、'.join(v)}", style="List Bullet")

    return doc


__all__ = ["CoursePlan", "CoursePlanStore", "build_course_plan", "rebuild"]
